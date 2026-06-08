"""
InsurRX 기획서 → .docx 변환 스크립트
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# 헬퍼 함수
# ─────────────────────────────────────────────
BLUE   = RGBColor(0x25, 0x63, 0xEB)   # #2563EB
TEAL   = RGBColor(0x0D, 0x94, 0x88)   # #0D9488
DARK   = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A
GRAY   = RGBColor(0x64, 0x74, 0x8B)   # #64748B
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
LIGHT  = RGBColor(0xF1, 0xF5, 0xF9)   # 배경 회색


def set_cell_bg(cell, hex_color: str):
    """테이블 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """셀 테두리 설정 (top/bottom/left/right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), kwargs[side].get('val', 'single'))
            border.set(qn('w:sz'), str(kwargs[side].get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), kwargs[side].get('color', 'auto'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if color:
            run.font.color.rgb = color
        run.font.bold = True
    return p


def add_body(doc, text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_code_block(doc, text):
    """코드/다이어그램 블록 — 회색 배경, Consolas 폰트."""
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        # 배경색 (단락 shading)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F1F5F9')
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def add_table(doc, headers, rows, header_bg='2563EB'):
    """헤더 강조 + 줄무늬 테이블 생성."""
    cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F8FAFC'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            set_cell_bg(row_cells[c_idx], bg)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────
# 문서 생성
# ─────────────────────────────────────────────
doc = Document()

# 페이지 여백
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 기본 폰트
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

# ──────────────────────────────────────────────────────────
# 표지
# ──────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('InsurRX')
r.font.name = '맑은 고딕'
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('서비스 기획서')
r2.font.size = Pt(20)
r2.font.color.rgb = DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('처방전 사진 한 장으로 내 보험 보상을 3초 안에 확인하는 AI 챗봇')
r3.font.size = Pt(12)
r3.font.color.rgb = GRAY
r3.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

# 표지 정보 테이블
info_table = doc.add_table(rows=3, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('버전', 'v1.0'),
    ('작성일', '2026-05-28'),
    ('분류', '내부 기획 문서'),
]
for i, (k, v) in enumerate(info_data):
    c0 = info_table.rows[i].cells[0]
    c1 = info_table.rows[i].cells[1]
    c0.text = k
    c1.text = v
    set_cell_bg(c0, 'EFF6FF')
    for para in c0.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(10)
    for para in c1.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)

doc.add_page_break()

# ──────────────────────────────────────────────────────────
# 목차 (수동)
# ──────────────────────────────────────────────────────────
add_heading(doc, '목차', level=1, color=BLUE)
toc_items = [
    ('1', '서비스 개요'),
    ('2', '문제 정의'),
    ('3', '솔루션'),
    ('4', '핵심 기능'),
    ('5', '기술 아키텍처'),
    ('6', '사용자 시나리오'),
    ('7', '타겟 사용자'),
    ('8', '비즈니스 모델'),
    ('9', '경쟁 분석'),
    ('10', '개발 현황'),
    ('11', '외부 서비스 구성'),
    ('12', '향후 로드맵'),
]
for num, title in toc_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{num}. {title}')
    run.font.size = Pt(11)

doc.add_page_break()

# ──────────────────────────────────────────────────────────
# 1. 서비스 개요
# ──────────────────────────────────────────────────────────
add_heading(doc, '1. 서비스 개요', level=1, color=BLUE)

add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ('서비스명',     'InsurRX (인슈알엑스)'),
        ('한 줄 정의',   '처방전 사진 한 장으로 내 보험 보상을 3초 안에 확인하는 AI 챗봇'),
        ('핵심 가치',    '미청구 보험금 자동 발굴 — 귀찮아서 포기했던 소액부터 찾아드립니다'),
        ('서비스 URL',  'https://songgwangsun.github.io/InsurRX/'),
        ('API URL',     'https://magnificent-celebration-production-2dd9.up.railway.app'),
        ('개발 과정',    '2026ICT AI 서비스 개발 과정'),
    ]
)

# ──────────────────────────────────────────────────────────
# 2. 문제 정의
# ──────────────────────────────────────────────────────────
add_heading(doc, '2. 문제 정의', level=1, color=BLUE)
add_heading(doc, '2-1. 시장 현황', level=2, color=TEAL)

add_body(doc, '국내 민간 보험 가입자는 전체 인구의 약 75% 수준이며, 1인당 평균 3.2개의 보험을 보유하고 있습니다. 그럼에도 실제 보험금 청구율은 현저히 낮습니다.')

p_quote = doc.add_paragraph()
r_q = p_quote.add_run('  연간 미청구 보험금 추정액: 2조 원 이상  (금융감독원 보험 소비자 실태조사 참고)')
r_q.font.bold = True
r_q.font.color.rgb = BLUE
r_q.font.size = Pt(11)
p_quote.paragraph_format.left_indent = Cm(1)

add_heading(doc, '2-2. 미청구의 3가지 원인', level=2, color=TEAL)
add_code_block(doc, '① 복잡성  — 수백 페이지 약관, 어려운 법률 용어\n② 귀찮음  — 소액(1~5만원)이라 청구 절차가 더 번거로움\n③ 무지    — 특약 존재 자체를 모르거나 보장 범위 불명확')

add_heading(doc, '2-3. 기존 방식의 한계', level=2, color=TEAL)
add_table(doc,
    headers=['기존 방법', '문제점'],
    rows=[
        ('보험사 앱',     '내 보험만 확인 가능, 비교 불가'),
        ('콜센터 상담',   '평균 대기 30분 이상'),
        ('보험 설계사',   '영업 목적 개입, 신뢰도 낮음'),
        ('직접 약관 검색','수백 페이지 중 해당 조항 찾기 불가'),
    ]
)

# ──────────────────────────────────────────────────────────
# 3. 솔루션
# ──────────────────────────────────────────────────────────
add_heading(doc, '3. 솔루션', level=1, color=BLUE)
add_heading(doc, 'InsurRX의 접근 방식', level=2, color=TEAL)
add_code_block(doc,
    '처방전/영수증 사진 1장\n'
    '        ↓\n'
    'GPT-4o Vision OCR (한글·영문·숫자 자동 인식)\n'
    '        ↓\n'
    '상병코드 · 약품명 · 금액 구조화\n'
    '        ↓\n'
    'Pinecone 벡터 DB 약관 검색 (의미 기반 RAG, 1,358개 조항)\n'
    '        ↓\n'
    'GPT-4o 보상 분석\n'
    '        ↓\n'
    '청구 가능 여부 + 예상 금액 + 약관 근거  [3초 이내]'
)

add_heading(doc, '핵심 차별점', level=2, color=TEAL)
add_table(doc,
    headers=['항목', 'InsurRX', '기존 서비스'],
    rows=[
        ('분석 단위',   '질병코드·약품 단위 정밀 분석', '진단명 키워드 매칭'),
        ('면책 조항',   'AI가 비급여·선천성 질환 사전 필터링', '수동 확인 필요'),
        ('약관 근거',   '적용된 조항 원문 인용', '결과만 제공'),
        ('다보험 비교', '보유 보험 전체 동시 매칭', '단일 보험사'),
        ('속도',        '3초 이내 응답', '수 분~수십 분'),
    ]
)

# ──────────────────────────────────────────────────────────
# 4. 핵심 기능
# ──────────────────────────────────────────────────────────
add_heading(doc, '4. 핵심 기능', level=1, color=BLUE)
add_heading(doc, '4-1. Vision OCR', level=2, color=TEAL)
add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ('엔진',       'GPT-4o Vision (gpt-4o-mini)'),
        ('인식 대상',  '진료비 영수증, 처방전, 약품 봉투'),
        ('추출 항목',  '병원명, 진료과, 상병코드(ICD), 처방일, 본인부담금, 약품명·용량·일수, 비급여 항목'),
        ('지원 포맷',  'JPEG, PNG, WEBP, HEIC(자동 변환)'),
        ('파싱 방식',  '정규식 1차 추출 → 누락 시 GPT-4o-mini LLM 보완'),
    ]
)

add_heading(doc, '4-2. RAG 보상 분석', level=2, color=TEAL)
add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ('Vector DB',  'Pinecone (AWS us-east-1, 1,358개 벡터)'),
        ('임베딩',     'OpenAI text-embedding-3-small (dim=1536)'),
        ('LLM',        'GPT-4o (보상 판단·금액 계산·약관 해석)'),
        ('출력',       '청구 가능 여부, 예상 보상 금액, 계산 내역, 약관 조항 원문 인용, AI 신뢰도'),
    ]
)

add_heading(doc, '보유 약관 4종', level=3, color=DARK)
add_table(doc,
    headers=['약관', '보험사', '종류', '벡터 수'],
    rows=[
        ('Hi2204',      '현대해상', '실손의료보험', '약 350개'),
        ('Hi2401',      '현대해상', '암보험 정액형', '약 480개'),
        ('My아이플러스', '삼성화재', '어린이보험', '약 390개'),
        ('갱신형 2501', 'AXA',     '치아보험', '약 138개'),
    ]
)

add_heading(doc, '4-3. Waitlist 이메일 수집', level=2, color=TEAL)
for item in [
    'Railway PostgreSQL 저장 (중복 방지)',
    'Formspree 폴백 (API 장애 시 자동 전환)',
    '선착순 1,000명 → 베타 출시 알림',
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

add_heading(doc, '4-4. 카카오톡 채널', level=2, color=TEAL)
add_body(doc, '채널 URL: https://pf.kakao.com/_xmxfjxjX\n향후 카카오 챗봇 시나리오 연동 예정')

# ──────────────────────────────────────────────────────────
# 5. 기술 아키텍처
# ──────────────────────────────────────────────────────────
add_heading(doc, '5. 기술 아키텍처', level=1, color=BLUE)
add_code_block(doc,
    '사용자 (브라우저/모바일)\n'
    '        │ HTTPS\n'
    '  GitHub Pages (index.html — 랜딩 + OCR UI)\n'
    '        │ POST /api/v1/...\n'
    '  Railway FastAPI (Python 3.12 / Uvicorn)\n'
    '    ├── /upload   → GPT-4o Vision OCR\n'
    '    ├── /analyze  → RAG Chain + GPT-4o\n'
    '    ├── /waitlist → PostgreSQL\n'
    '    └── /result   → PostgreSQL\n'
    '        │\n'
    '  ├── Pinecone Vector DB (1,358 vectors, 4개 약관)\n'
    '  └── Railway PostgreSQL (Waitlist, 분석결과)'
)

add_heading(doc, '기술 스택 상세', level=2, color=TEAL)
add_table(doc,
    headers=['계층', '기술'],
    rows=[
        ('Frontend',    'HTML5, CSS3, Vanilla JS'),
        ('Backend',     'FastAPI 0.115, Python 3.12, Uvicorn'),
        ('ORM',         'SQLAlchemy 2.0 Async'),
        ('DB',          'PostgreSQL (운영), SQLite (로컬)'),
        ('AI/OCR',      'OpenAI gpt-4o-mini (Vision)'),
        ('AI/LLM',      'OpenAI gpt-4o (RAG 분석)'),
        ('임베딩',       'OpenAI text-embedding-3-small'),
        ('Vector DB',   'Pinecone (Starter, AWS us-east-1)'),
        ('RAG 프레임워크','LangChain 0.3.25'),
        ('인프라',       'Railway (백엔드), GitHub Pages (프론트)'),
        ('CI/CD',       'GitHub Actions'),
        ('컨테이너',     'Docker (Railway Dockerfile 배포)'),
    ]
)

# ──────────────────────────────────────────────────────────
# 6. 사용자 시나리오
# ──────────────────────────────────────────────────────────
add_heading(doc, '6. 사용자 시나리오', level=1, color=BLUE)
add_heading(doc, '메인 플로우 (7단계)', level=2, color=TEAL)
add_code_block(doc,
    'STEP 1  사진 입력 방법 선택\n'
    '        [📷 카메라 촬영] 또는 [🖼️ 갤러리 선택]\n'
    '\n'
    'STEP 2  이미지 미리보기 확인\n'
    '        → "AI 분석 시작" 클릭\n'
    '\n'
    'STEP 3  OCR 분석 중 (평균 3~5초)\n'
    '        GPT-4o Vision이 처방전 텍스트 추출\n'
    '\n'
    'STEP 4  인식 결과 확인\n'
    '        병원명 · 상병코드 · 약품 · 본인부담금 표시\n'
    '        → "보상 분석 받기" 클릭\n'
    '\n'
    'STEP 5  보험 선택\n'
    '        가입 보험 체크박스 선택 (복수 선택 가능)\n'
    '        → "분석 시작" 클릭\n'
    '\n'
    'STEP 6  RAG 분석 중 (평균 5~10초)\n'
    '        Pinecone 벡터 검색 + GPT-4o 보상 판단\n'
    '\n'
    'STEP 7  분석 결과 리포트\n'
    '        · ✅ 청구 가능 / ❌ 청구 어려움\n'
    '        · 예상 보상 금액 (₩ XX,XXX)\n'
    '        · 계산 내역 (본인부담금 - 공제액)\n'
    '        · 적용 약관 조항 + 원문 인용\n'
    '        · AI 신뢰도 %'
)

add_heading(doc, '엔드투엔드 테스트 결과 (2026-05-26)', level=2, color=TEAL)
add_code_block(doc,
    '입력:  급성상기도감염(J06.9), 28,500원, 현대해상 실손 선택\n'
    '출력:  is_claimable=true\n'
    '       estimated_payout=18,500원\n'
    '       공제액: 통원 의원급 10,000원\n'
    '       confidence=0.9'
)

# ──────────────────────────────────────────────────────────
# 7. 타겟 사용자
# ──────────────────────────────────────────────────────────
add_heading(doc, '7. 타겟 사용자', level=1, color=BLUE)
add_heading(doc, '메인 페르소나', level=2, color=TEAL)
add_code_block(doc,
    '"37세 IT 기업 과장 워킹맘"\n'
    '\n'
    '- 아이 둘, 월 2~3회 소아과·이비인후과 방문\n'
    '- 실손보험 가입 중이지만 소액 귀찮아서 청구 포기\n'
    '- "병원 영수증이 쌓이는데 청구하기 너무 복잡해요"\n'
    '- 카카오톡 주 사용, 모바일 우선'
)

add_heading(doc, '세컨더리 타겟', level=2, color=TEAL)
add_table(doc,
    headers=['세그먼트', '특징'],
    rows=[
        ('만성질환자 (40~60대)', '월 1~2회 병원 방문, 보험 청구 의향 높음'),
        ('부모 (자녀 보험 보유)', '어린이보험 특약 활용 빈도 높음'),
        ('보험 설계사', '고객 상담 도구로 활용 (B2B)'),
    ]
)

# ──────────────────────────────────────────────────────────
# 8. 비즈니스 모델
# ──────────────────────────────────────────────────────────
add_heading(doc, '8. 비즈니스 모델', level=1, color=BLUE)
add_heading(doc, 'B2C Freemium', level=2, color=TEAL)
add_table(doc,
    headers=['플랜', '가격', '주요 기능'],
    rows=[
        ('베이직 (무료)', '₩0',         '월 5회 보상 확인, OCR, 기본 약관 매칭'),
        ('라이트',        '₩2,900/월',  '무제한 확인, 면책 필터링, 원클릭 청구 대행, 가족 3인, 12개월 이력'),
    ]
)
add_body(doc, '  베타 기간: 라이트 플랜 선착순 1,000명 무료 체험', bold=True, color=BLUE)

add_heading(doc, 'B2B (중장기)', level=2, color=TEAL)
add_table(doc,
    headers=['채널', '내용'],
    rows=[
        ('보험 GA SaaS', '설계사 대상 API 제공 (건당 과금)'),
        ('제휴 보험사',   '미청구 보험금 안내 캠페인 (CPA 수수료)'),
        ('병원·약국',    'POS 연동 영수증 자동 분석'),
    ]
)

# ──────────────────────────────────────────────────────────
# 9. 경쟁 분석
# ──────────────────────────────────────────────────────────
add_heading(doc, '9. 경쟁 분석', level=1, color=BLUE)
add_table(doc,
    headers=['서비스', '강점', '한계', 'InsurRX 대비'],
    rows=[
        ('토스 보험',       '대형 플랫폼, 신뢰도', '단순 보험료 비교, OCR 없음', '청구 분석 깊이'),
        ('굿닥',            '병원 정보 강점', '보험 분석 기능 없음', '보험 도메인 전문성'),
        ('시그널플래너',    '보험 포트폴리오 관리', '수동 입력, 실시간 분석 없음', '자동 OCR 연결'),
        ('카카오페이 보험', '간편 가입/청구', '자사 제휴 보험만 지원', '다사 약관 비교'),
    ]
)
add_body(doc, 'InsurRX 핵심 차별점: 질병코드·약품 단위 면책 조항 AI 필터링', bold=True, color=BLUE)

# ──────────────────────────────────────────────────────────
# 10. 개발 현황
# ──────────────────────────────────────────────────────────
add_heading(doc, '10. 개발 현황', level=1, color=BLUE)
add_heading(doc, '✅ 완료 (2026-06-08 기준)', level=2, color=TEAL)

completed = [
    ('인프라',
     ['FastAPI 백엔드 Railway 배포 · PostgreSQL 연동',
      'GitHub Pages 3페이지 자동 배포 (index / dashboard / admin)',
      'CI(pytest 55개) · CD(Railway + GitHub Pages) 파이프라인',
      'CORS 설정, Docker 컨테이너 배포']),
    ('인증 · 회원',
     ['JWT 회원가입 · 로그인 (bcrypt + python-jose)',
      '프로필 수정 · 비밀번호 변경 API',
      '관리자 권한 분리 (is_admin) · 회원 활성/비활성 토글']),
    ('AI 파이프라인',
     ['GPT-4o Vision OCR — 처방전 인식 실서비스 확인',
      'Pinecone 벡터 DB — 4개 공개 약관 1,358개 벡터 업서트 완료',
      '개인 보험증권 업로드 → 청킹 → 임베딩 → 개인 네임스페이스(user-{id}) 업서트',
      'RAG 보상 분석 — 공개 약관 + 개인 보험증권 통합 검색 (JWT 자동 연동)',
      '엔드투엔드 테스트 통과 (신뢰도 92%)']),
    ('API (총 20+ 엔드포인트)',
     ['POST /upload/ — GPT-4o Vision OCR · 파싱(정규식+LLM 폴백)',
      'POST /analyze/ — 통합 RAG 보상 분석 (JWT 선택 인증)',
      'POST /auth/register · /auth/login · PATCH /auth/me · POST /auth/me/change-password',
      'POST /my/policies/upload · GET/DELETE /my/policies/ — 개인 보험증권 관리',
      'GET/POST/PATCH/DELETE /insurers/{id}/products/ — 보험사·상품 CRUD',
      'GET /my/analyses/ · DELETE /my/analyses/{id} — 분석 히스토리',
      'GET /admin/users · PATCH /admin/users/{id}/toggle-* · GET /admin/stats',
      'POST /kakao/skill — 카카오 i 오픈빌더 스킬 서버']),
    ('프론트엔드',
     ['index.html — 랜딩페이지 (히어로, How it works, OCR 데모, 요금제, Waitlist)',
      'dashboard.html — 회원 대시보드 (보상 분석 5단계 · 보험증권 관리 · 히스토리 · 계정 설정)',
      'admin.html — 관리자 패널 (보험사·상품 CRUD · 회원 목록 · 시스템 통계)']),
    ('카카오 챗봇',
     ['카카오 i 오픈빌더 스킬 서버 구현 (POST /kakao/skill)',
      '대화 세션 관리 (KakaoSession DB) — 이미지 수신 → 약관 선택 → 비동기 분석 → 결과 확인',
      '카카오 5초 응답 제한 대응: OCR+RAG 백그라운드 처리 후 결과 폴링',
      '랜딩페이지 히어로 CTA: 카카오톡 채팅 직링크 연결']),
    ('테스트',
     ['단위·통합 테스트 55개 통과 (auth 6 · insurer 6 · user_policy 6 · kakao 7 + 기존 30)']),
]
for category, items in completed:
    p = doc.add_paragraph()
    r = p.add_run(f'■ {category}')
    r.font.bold = True
    r.font.size = Pt(10)
    for item in items:
        p2 = doc.add_paragraph(style='List Bullet')
        p2.add_run(item).font.size = Pt(9.5)
        p2.paragraph_format.left_indent = Cm(0.5)

add_heading(doc, '🔲 향후 작업', level=2, color=TEAL)
add_table(doc,
    headers=['우선순위', '기능'],
    rows=[
        ('🔴 높음', 'CBT — Waitlist 유저 초대 · 쿼리 로그 분석 · 응답 정확도 개선'),
        ('🟡 중간', '마이데이터 API — 내 보험 자동 불러오기 (금융위 오픈API)'),
        ('🟡 중간', '원클릭 청구 대행 — 보험사 API 연동 (파일럿)'),
        ('🟢 낮음', '가족 계정 통합 관리'),
        ('🟢 낮음', '청구 이력 12개월 보관 · 알림 서비스'),
    ]
)

# ──────────────────────────────────────────────────────────
# 11. 외부 서비스 구성
# ──────────────────────────────────────────────────────────
add_heading(doc, '11. 외부 서비스 구성', level=1, color=BLUE)
add_table(doc,
    headers=['서비스', '용도', '상태'],
    rows=[
        ('OpenAI (gpt-4o-mini)',           'OCR 텍스트 추출',                  '✅ 운영 중'),
        ('OpenAI (gpt-4o)',                '보상 분석 LLM',                    '✅ 운영 중'),
        ('OpenAI (text-embedding-3-small)','공개 약관 + 개인 보험증권 임베딩',  '✅ 운영 중'),
        ('Pinecone (insurrx-policies)',    '공개 약관 + 개인 보험증권 벡터 검색 (네임스페이스 분리)', '✅ 운영 중'),
        ('Railway (iad 리전)',              'FastAPI 백엔드 + PostgreSQL 호스팅','✅ 운영 중'),
        ('GitHub Pages',                   'index / dashboard / admin 호스팅',  '✅ 운영 중'),
        ('Formspree',                      'Waitlist 이메일 수집 폴백',         '✅ 운영 중'),
        ('KakaoTalk 채널',                 '챗봇 채널 (오픈빌더 스킬 서버 연결)','✅ 채널 개설'),
    ]
)

add_heading(doc, '보안 및 컴플라이언스', level=2, color=TEAL)
security_items = [
    '업로드 이미지: OCR 텍스트 추출 후 메모리에서 즉시 해제 (서버 저장 없음)',
    '개인정보: 이름·주민번호 등 식별 정보 비식별화 처리',
    '전송 구간: HTTPS(TLS 1.3) 암호화',
    '데이터 보유: 분석 결과만 session_id로 보관, 원본 이미지 미보관',
    '준수 기준: 개인정보보호법, K-HIPAA 가이드라인',
]
for item in security_items:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item).font.size = Pt(10)

# ──────────────────────────────────────────────────────────
# 12. 향후 로드맵
# ──────────────────────────────────────────────────────────
add_heading(doc, '12. 향후 로드맵', level=1, color=BLUE)
add_table(doc,
    headers=['시기', '주요 마일스톤'],
    rows=[
        ('2026 Q2 (현재)', '✅ MVP 개발 완료 · 실서비스 배포 · OCR+RAG 동작 확인 · CBT 시작'),
        ('2026 Q3',        '카카오톡 챗봇 정식 운영 · 마이데이터 API 연동 · 원클릭 청구 (파일럿) · 약관 확대'),
        ('2026 Q4',        '라이트 플랜 유료 전환 · B2B GA SaaS API 베타 · 가족 계정·이력 관리'),
        ('2027 Q1~',       '마이데이터 기반 보험 자동 매칭 · 원클릭 청구 전사 확대 · 카카오페이·토스 제휴'),
    ]
)

# ──────────────────────────────────────────────────────────
# 부록 — API 명세
# ──────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '부록 — API 명세 요약', level=1, color=BLUE)

add_heading(doc, 'POST /api/v1/upload/', level=2, color=TEAL)
add_code_block(doc,
    '// Request: multipart/form-data\n'
    '{ "file": "<이미지 파일>" }\n'
    '\n'
    '// Response 200\n'
    '{\n'
    '  "filename": "prescription.jpg",\n'
    '  "ocr_raw": "○○내과의원\\n내과\\nJ06.9 급성상기도감염...",\n'
    '  "parsed": {\n'
    '    "hospital": "○○내과의원",\n'
    '    "department": "내과",\n'
    '    "icd_code": "J06.9",\n'
    '    "diagnosis": "급성상기도감염",\n'
    '    "prescription_date": "2026-05-28",\n'
    '    "total_amount": 28500,\n'
    '    "drugs": [\n'
    '      { "name": "아목시실린 500mg", "dosage": "500mg", "days": 3, "is_nonbenefit": false }\n'
    '    ]\n'
    '  }\n'
    '}'
)

add_heading(doc, 'POST /api/v1/analyze/', level=2, color=TEAL)
add_code_block(doc,
    '// Request  (Authorization: Bearer <token> 선택)\n'
    '{\n'
    '  "parsed": { /* ParsedDocument */ },\n'
    '  "policy_ids": ["hyundai-silson-v4"]  // 선택사항 — 로그인 시 개인 보험증권 자동 포함\n'
    '}\n'
    '\n'
    '// Response 200\n'
    '{\n'
    '  "session_id": "uuid",\n'
    '  "is_claimable": true,\n'
    '  "estimated_payout": 18500,\n'
    '  "breakdown": { "본인부담금": 28500, "공제액(통원 의원급)": -10000 },\n'
    '  "coverage_items": [\n'
    '    {\n'
    '      "clause_id": "hyundai-silson-3-2",\n'
    '      "policy_name": "현대해상 실손의료비",\n'
    '      "article": "제3조 ②항",\n'
    '      "is_covered": true,\n'
    '      "reason": "통원 의원급 해당, 공제 후 보상",\n'
    '      "citation": "통원의료비는 1회당 의원 1만원을 공제한 금액을 보상한다."\n'
    '    }\n'
    '  ],\n'
    '  "confidence": 0.92,\n'
    '  "llm_summary": "실손의료비 통원 항목으로 약 ₩18,500 청구 가능합니다."\n'
    '}'
)

add_heading(doc, 'POST /api/v1/auth/register  ·  /auth/login', level=2, color=TEAL)
add_code_block(doc,
    '// POST /auth/register\n'
    '{ "name": "홍길동", "email": "user@example.com", "password": "pass1234" }\n'
    '→ 201 { "access_token": "eyJ...", "token_type": "bearer", "user": { ... } }\n'
    '\n'
    '// POST /auth/login  (OAuth2PasswordBearer)\n'
    'username=user@example.com&password=pass1234\n'
    '→ 200 { "access_token": "eyJ...", "user": { "id": 1, "is_admin": false, ... } }'
)

add_heading(doc, 'POST /api/v1/my/policies/upload', level=2, color=TEAL)
add_code_block(doc,
    '// Request: multipart/form-data  (Authorization 필수)\n'
    '{ "file": "<PDF or 이미지>", "product_id": 2 }  // product_id 선택\n'
    '\n'
    '// Response 202 — 임베딩은 백그라운드 처리\n'
    '{ "policy_id": 1, "original_name": "내보험증권.pdf",\n'
    '  "status": "pending", "message": "업로드 완료. 임베딩 진행 중..." }'
)

add_heading(doc, 'POST /api/v1/kakao/skill', level=2, color=TEAL)
add_code_block(doc,
    '// 카카오 i 오픈빌더 스킬 서버 요청 (v2)\n'
    '{\n'
    '  "userRequest": { "utterance": "보험금 분석", "user": { "id": "kakao_uid" } },\n'
    '  "action": { "params": { "image": { "url": "https://..." } } }\n'
    '}\n'
    '\n'
    '// Response — 카드 형식\n'
    '{\n'
    '  "version": "2.0",\n'
    '  "template": {\n'
    '    "outputs": [{ "basicCard": { "title": "✅ 보험금 청구 가능 — ₩18,500", ... } }],\n'
    '    "quickReplies": [{ "label": "🔄 다시 분석", "action": "message" }]\n'
    '  }\n'
    '}'
)

# ──────────────────────────────────────────────────────────
# 푸터 (마지막 단락)
# ──────────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run('본 문서는 InsurRX 서비스의 내부 기획 문서입니다.  |  문의: songgs33@gmail.com')
r_f.font.size = Pt(9)
r_f.font.color.rgb = GRAY
r_f.font.italic = True

# ──────────────────────────────────────────────────────────
# 저장
# ──────────────────────────────────────────────────────────
from pathlib import Path
OUTPUT = Path(__file__).parent / '기획서_InsurRX_v1.0.docx'
doc.save(OUTPUT)
print(f'저장 완료: {OUTPUT}')
